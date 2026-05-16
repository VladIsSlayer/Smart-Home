#!/usr/bin/env python3
"""Общий docx: пошаговая инструкция по отчётам и скриншотам для ЛР4–ЛР9."""

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]

LABS = [
    {
        "num": 4,
        "title": "Передача управляющих команд",
        "run": "LR4",
        "focus": "Отдельные маршруты /control_* для управления устройствами.",
        "steps": [
            "Запустить: cd LR4, py -3 app.py.",
            "Открыть http://127.0.0.1:5000/index.html, войти как администратор (PIN 1234).",
            "На admin.html: включить чайник, открыть/закрыть шторы, применить RGB ламп.",
            "Открыть F12 → Network, повторить команду — увидеть запрос control_*.",
        ],
        "photos": [
            ("Рис. 1", "Панель администратора с блоками устройств после выполнения команд."),
            ("Рис. 2", "Сообщение внизу страницы («Чайник…», «Шторы…») или вкладка Network с ответом JSON ok: true."),
            ("Рис. 3", "Консоль терминала Flask с логами [SmartKettle.control], [SmartCurtains.control]."),
            ("Рис. 4", "Страница пользователя (user.html) — ограниченный набор кнопок."),
        ],
        "code": ["LR4/things.py — методы control()", "LR4/app.py — маршруты /control_*"],
    },
    {
        "num": 5,
        "title": "Валидация входящих данных",
        "run": "LR5",
        "focus": "Проверка чисел (try/except) и строк (regex) на сервере.",
        "steps": [
            "Запустить LR5, открыть admin.html.",
            "В адресной строке вызвать ошибку: .../control_smart_kettle?action=on&target_temp=abc",
            "Вызвать успех: .../control_smart_kettle?action=on&target_temp=85",
            "Пылесос: .../control_robot_vacuum?action=start&mode=wrong — ошибка режима.",
        ],
        "photos": [
            ("Рис. 1", "Браузер с JSON ok: false и текстом ошибки (неверное число)."),
            ("Рис. 2", "Браузер с JSON ok: true при корректных параметрах."),
            ("Рис. 3", "Ошибка режима пылесоса (eco/auto/turbo)."),
            ("Рис. 4", "Фрагмент things.py — _parse_int или VACUUM_MODE_PATTERN."),
        ],
        "code": ["LR5/things.py — валидация в control()"],
    },
    {
        "num": 6,
        "title": "Автоматические режимы",
        "run": "LR6",
        "focus": "Класс HomeAutomation — климат влияет на лампы и шторы.",
        "steps": [
            "Запустить LR6, admin.html.",
            "Климат: текущая T ниже цели на 2+ °C (например цель 26 °C).",
            "Подождать 10–20 с — появится сообщение «Автоматика: …».",
            "Проверить изменение яркости ламп или положения штор.",
        ],
        "photos": [
            ("Рис. 1", "Блок климата: текущая и целевая температура до автоматики."),
            ("Рис. 2", "Сообщение «Автоматика: …» на панели администратора."),
            ("Рис. 3", "Состояние ламп/штор после срабатывания правил."),
            ("Рис. 4", "Лог Flask: [HomeAutomation.run_after_sensor_update]."),
        ],
        "code": ["LR6/things.py — class HomeAutomation"],
    },
    {
        "num": 7,
        "title": "Сбор и хранение данных",
        "run": "LR7",
        "focus": "Logger — insert_temperature, insert_humidity, MongoDB или JSON.",
        "steps": [
            "Запустить LR7, держать admin.html открытой 30–40 с.",
            "Открыть LR7/data/IOT_logger_db.json (или MongoDB Compass).",
            "Убедиться, что есть массивы Temperature и Humidity с timeStamp.",
        ],
        "photos": [
            ("Рис. 1", "Файл LR7/data/IOT_logger_db.json с записями температуры и влажности."),
            ("Рис. 2", "Лог Flask: [Logger.insert_temperature] / insert_humidity."),
            ("Рис. 3", "MongoDB Compass — коллекции Temperature, Humidity (если MongoDB установлен)."),
            ("Рис. 4", "Фрагмент class Logger в things.py."),
        ],
        "code": ["LR7/things.py — class Logger", "LR7/app.py — вызов insert_* в connect_*"],
    },
    {
        "num": 8,
        "title": "Анализ данных",
        "run": "LR8",
        "focus": "Среднее и максимум T/H, блок «Анализ данных», /api/analysis.",
        "steps": [
            "Запустить LR8, admin.html 30 с (накопление лога).",
            "Задать цель климата 27 °C — «Применить цели климата».",
            "Проверить блок «Анализ данных» — цифры не прочерки.",
            "Открыть http://127.0.0.1:5000/api/analysis в браузере.",
        ],
        "photos": [
            ("Рис. 1", "Блок «Анализ данных» со средней/макс. температурой и влажностью."),
            ("Рис. 2", "Страница /api/analysis с JSON analysis."),
            ("Рис. 3", "LR8/data/IOT_logger_db.json — несколько записей Temperature."),
            ("Рис. 4", "Климат: текущая T растёт к цели (для связи с логом)."),
        ],
        "code": ["LR8/things.py — get_average_*, get_max_*", "LR8/app.py — /api/analysis"],
    },
    {
        "num": 9,
        "title": "Визуализация (Chart.js)",
        "run": "LR9",
        "focus": "График температуры, полный функционал ЛР4–8 + сценарии.",
        "steps": [
            "Запустить LR9, admin.html 40–60 с.",
            "Применить цель климата 28 °C, дождаться графика.",
            "Шторы: слайдер 75 % → «Сохранить для Открыть» → «Закрыть» → «Открыть» (до 75 %).",
            "Сценарий «Теплый вечер» → Применить.",
            "Открыть /api/chart/temperature.",
        ],
        "photos": [
            ("Рис. 1", "Блок «Анализ и визуализация» с цифрами и линейным графиком Chart.js."),
            ("Рис. 2", "Шторы: сохранённое открытие 75 %, текущее положение при движении."),
            ("Рис. 3", "Список сценариев и результат «Сценарий применён»."),
            ("Рис. 4", "JSON /api/chart/temperature (labels, values) или Network."),
        ],
        "code": ["LR9/things.py — get_temperature_chart_data", "LR9/server/script.js — loadTemperatureChart"],
    },
]


def build_document() -> Document:
    doc = Document()
    t = doc.add_heading("Инструкция по отчётам: лабораторные 4–9", level=0)
    t.runs[0].font.size = Pt(18)
    doc.add_paragraph(
        "Умный дом — единый проект. Для отчёта по каждой лаборатории запускайте только папку LR{N} "
        "(порт 5000). Скриншоты делайте с подписями «Рис. N» и кратким пояснением под каждым."
    )
    doc.add_paragraph("Общие требования к отчёту: титул, цель, задание, код, скриншоты (не менее 2), выводы.")
    doc.add_paragraph("")

    doc.add_heading("Общая подготовка", level=1)
    doc.add_paragraph("pip install -r requirements.txt", style="List Bullet")
    doc.add_paragraph("Один терминал: cd LR{N} && py -3 app.py", style="List Bullet")
    doc.add_paragraph("Браузер: http://127.0.0.1:5000/index.html (не file://)", style="List Bullet")
    doc.add_paragraph("Админ: PIN 1234", style="List Bullet")
    doc.add_paragraph(
        "Для ЛР7–9 данные пишутся в LR{N}/data/IOT_logger_db.json, если MongoDB не запущен.",
        style="List Bullet",
    )

    for lab in LABS:
        doc.add_page_break()
        doc.add_heading(f"Лабораторная работа № {lab['num']}: {lab['title']}", level=1)
        doc.add_paragraph(f"Папка запуска: {lab['run']}/")
        doc.add_paragraph(f"Акцент отчёта: {lab['focus']}")

        doc.add_heading("Порядок действий", level=2)
        for i, step in enumerate(lab["steps"], 1):
            doc.add_paragraph(f"{i}. {step}", style="List Number")

        doc.add_heading("Скриншоты для отчёта (минимум 2, рекомендуется 4)", level=2)
        for title, desc in lab["photos"]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"{title}. ").bold = True
            p.add_run(desc)

        doc.add_heading("Что вставить из кода", level=2)
        for item in lab["code"]:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("Пример вывода", level=2)
        doc.add_paragraph(
            f"В ЛР{lab['num']} реализовано: {lab['focus']} "
            "Интерфейс и сервер обмениваются JSON; управление и мониторинг разделены по маршрутам."
        )

    doc.add_page_break()
    doc.add_heading("Проверка LR9 как итоговой демонстрации", level=1)
    doc.add_paragraph(
        "Запуск LR9 позволяет одним сеансом снять доказательства для тем ЛР4–9: "
        "control, валидация, автоматика, лог, анализ, график, сценарии. "
        "В отчёте по каждой ЛР укажите, что использовалась соответствующая папка LR{N}, "
        "а скриншоты — с пометкой, какой пункт лабораторной они подтверждают."
    )
    return doc


def main() -> None:
    out = ROOT / "Otchet_LR4-9_instrukciya_skrinshoty.docx"
    build_document().save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
