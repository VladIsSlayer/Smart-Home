#!/usr/bin/env python3
"""Генерация docx-файлов с содержанием отчёта для ЛР4–ЛР9."""

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]

LABS = [
    {
        "num": 4,
        "title": "Организация передачи управляющих команд на удалённое оборудование",
        "goal": "Приобретение навыков организации передачи управляющих команд на удалённое оборудование.",
        "done": [
            "Маршруты /connect_* оставлены только для мониторинга (ЛР3).",
            "Добавлены маршруты /control_* и метод control() у каждой вещи.",
            "Кнопки интерфейса отправляют GET-запросы с параметрами команд.",
            "Ответ сервера отображается в сообщении на панели.",
        ],
        "screens": [
            "Страница администратора с кнопками управления.",
            "Страница пользователя (ограниченный набор команд).",
            "Консоль Flask с логом выполнения control().",
            "Пример успешной команды (вкл. чайник, шторы, лампы, пылесос).",
        ],
        "conclusions": [
            "Управление вынесено в отдельные запросы, мониторинг не изменён.",
            "Каждая вещь обрабатывает свой набор параметров.",
        ],
    },
    {
        "num": 5,
        "title": "Создание и настройка системы управления оборудованием",
        "goal": "Приобретение навыков проверки и приведения входящих данных к требуемому формату.",
        "done": [
            "Числовые поля проверяются через try/except и приведение типов.",
            "Строковые поля (режим пылесоса eco|auto|turbo, действие штор) — регулярными выражениями.",
            "При ошибке возвращается JSON с ok: false и текстом причины.",
        ],
        "screens": [
            "Успешная команда с корректными данными.",
            "Ошибка при передаче неверного числа (например, target_temp=abc).",
            "Ошибка при неверном режиме пылесоса.",
            "Лог сервера с сообщением об отклонении данных.",
        ],
        "conclusions": [
            "Валидация на стороне сервера предотвращает запись некорректных значений в объекты вещей.",
            "Использованы два формата: число и строка с regex.",
        ],
    },
    {
        "num": 6,
        "title": "Реализация полуавтоматических и автоматических режимов управления",
        "goal": "Наладить обмен данными между устройствами и автоматическое переключение состояний.",
        "done": [
            "Класс HomeAutomation связывает датчики температуры/влажности с лампами и шторами.",
            "При низкой температуре: повышается яркость, шторы закрываются.",
            "При высокой влажности: шторы открываются.",
            "Сообщения автоматики передаются в JSON и отображаются в интерфейсе.",
        ],
        "screens": [
            "Панель до срабатывания автоматики.",
            "Панель после срабатывания (изменились шторы/лампы).",
            "Лог сервера с записями HomeAutomation.",
        ],
        "conclusions": [
            "Автоматика реализует связи, заложенные в проекте ЛР1 (климат влияет на комфорт дома).",
        ],
    },
    {
        "num": 7,
        "title": "Создание и настройка системы сбора данных",
        "goal": "Приобретение навыков долгосрочного хранения данных системы.",
        "done": [
            "Класс Logger: insert_temperature, insert_humidity.",
            "Хранение в MongoDB (localhost) или в JSON-файле data/IOT_logger_db.json.",
            "Дублирующие подряд значения не записываются.",
        ],
        "screens": [
            "Содержимое коллекции Temperature (MongoDB Compass или mongo shell).",
            "Содержимое коллекции Humidity.",
            "Файл data/IOT_logger_db.json (если MongoDB не установлен).",
            "Лог сервера при insert_temperature / insert_humidity.",
        ],
        "conclusions": [
            "Система накапливает историю показаний для последующего анализа.",
        ],
    },
    {
        "num": 8,
        "title": "Модуль анализа данных в системах интернета вещей",
        "goal": "Приобретение навыков анализа накопленных данных.",
        "done": [
            "Среднее и максимальное значение температуры.",
            "Среднее и максимальное значение влажности.",
            "Блок «Анализ данных» на странице администратора, API /api/analysis.",
        ],
        "screens": [
            "Блок со статистикой на admin.html.",
            "Ответ /api/analysis в браузере или Postman.",
            "Несколько записей в БД перед расчётом статистики.",
        ],
        "conclusions": [
            "Анализ выполняется по данным из Logger, без изменения логики сбора.",
        ],
    },
    {
        "num": 9,
        "title": "Настройка системы визуализации данных",
        "goal": "Приобретение навыков визуализации данных для управления смарт-устройствами.",
        "done": [
            "Метод get_temperature_chart_data и API /api/chart/temperature.",
            "Линейный график Chart.js на admin.html.",
            "Периодическое обновление графика при опросе датчиков.",
        ],
        "screens": [
            "График температуры на панели администратора.",
            "Ответ /api/chart/temperature (labels и values).",
            "График после накопления нескольких точек данных.",
        ],
        "conclusions": [
            "Визуализация упрощает контроль динамики температуры во времени.",
        ],
    },
]


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.runs[0].font.size = Pt(16)


def build_doc(lab: dict) -> Document:
    doc = Document()
    add_title(doc, f"Лабораторная работа № {lab['num']}")
    doc.add_paragraph(f"{lab['title']}")
    doc.add_paragraph("Умный дом — тема проекта")
    doc.add_paragraph("")

    doc.add_heading("Содержание отчёта", level=1)
    doc.add_paragraph("1. Титульный лист")
    doc.add_paragraph("2. Цель работы")
    doc.add_paragraph("3. Задание")
    doc.add_paragraph("4. Ход работы / реализация")
    doc.add_paragraph("5. Листинги (things.py, app.py, фрагменты script.js)")
    doc.add_paragraph("6. Скриншоты")
    doc.add_paragraph("7. Выводы")

    doc.add_heading("Цель работы", level=1)
    doc.add_paragraph(lab["goal"])

    doc.add_heading("Что реализовано в проекте", level=1)
    for item in lab["done"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Какие скриншоты включить", level=1)
    for i, item in enumerate(lab["screens"], 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    doc.add_heading("Выводы (пример формулировок)", level=1)
    for item in lab["conclusions"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Запуск", level=1)
    n = lab["num"]
    doc.add_paragraph(f"cd LR{n}")
    doc.add_paragraph("pip install flask pymongo python-docx")
    doc.add_paragraph("py -3 app.py")
    doc.add_paragraph("Браузер: http://127.0.0.1:5000/index.html")
    if n >= 7:
        doc.add_paragraph(
            "Для ЛР7–9: установите MongoDB или используйте JSON-файл data/IOT_logger_db.json "
            "(создаётся автоматически, если MongoDB недоступен)."
        )

    return doc


def main() -> None:
    for lab in LABS:
        folder = ROOT / f"LR{lab['num']}"
        folder.mkdir(parents=True, exist_ok=True)
        out = folder / f"Otchet_LR{lab['num']}_soderzhanie.docx"
        doc = build_doc(lab)
        doc.save(out)
        print(f"OK LR{lab['num']}")


if __name__ == "__main__":
    main()
